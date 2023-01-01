from gpt_index import Document, GPTListIndex
import gradio as gr
import openai
import os
import PyPDF2
import docx
import pytesseract
from PIL import Image

def pdftotext(file_name):
  """
  Function to extract text from .pdf format files
  """

  text = []
  # Open the PDF file in read-binary mode
  with open(file_name, 'rb') as file:
    # Create a PDF object
    pdf = PyPDF2.PdfReader(file)

    # Get the number of pages in the PDF document
    num_pages = len(pdf.pages)

    # Iterate over every page
    for page in range(num_pages):
      # Extract the text from the page
      result = pdf.pages[page].extract_text()
      text.append(result)

  text = "\n".join(text)

  return text

def docxtotext(file_name):
  """
  Function to read .docx format files
  """
  # Open the Word document
  document = docx.Document(file_name)

  # Extract the text from the document
  text = '\n'.join([paragraph.text for paragraph in document.paragraphs])

  return text

def readtextfile(file_name):
  """
  Function to read .txt format files
  """

  # Open the Text document
  with open(file_name, 'r') as file:
    text = file.read()

  return text

def imagetotext(file_name):
  """
  Function to extract text from images
  """
  # Open the image using PIL
  image = Image.open(file_name)

  # Extract the text from the image
  text = pytesseract.image_to_string(image)

  return text

def preprocesstext(text):
  """
  Function to preprocess text
  """
  # Split the string into lines
  lines = text.splitlines()
  # Use a list comprehension to filter out empty lines
  lines = [line for line in lines if line.strip()]
  # Join the modified lines back into a single string
  text = '\n'.join(lines)

  return text

def processfiles(files):
  """
  Function to extract text from documents
  """
  textlist = []

  # Iterate over provided files
  for file in files:
    # Get file name
    file_name = file.name
    # Get extention of file name
    ext = file_name.split(".")[-1].lower()

    # Process document based on extention
    if ext == "pdf":
      text = pdftotext(file_name)
    elif ext == "docx":
      text = docxtotext(file_name)
    elif ext == "txt":
      text = readtextfile(file_name)
    elif ext in ["png", "jpg", "jpeg"]:
      text = imagetotext(file_name)
    else:
      text = ""

    # Preprocess text
    text = preprocesstext(text)

    # Append the text to final result
    textlist.append(text)

  return textlist

def createdocuments(textlist):
  """
  Function to create documents as needed for indexing.
  """
  documents = []
  # Create Document for indexing
  for text in textlist:
    documents.append(Document(text))

  return documents

def fileformatvaliditycheck(files):
  """
  Function to check validity of file formats
  """

  for file1 in files:
    file_name = file1.name
    # Get extention of file name
    ext = file_name.split(".")[-1].lower()

    if ext not in ["pdf", "txt", "docx", "png", "jpg", "jpeg"]:
      return False
  return True

def openaiapikeyvaliditycheck(openaikey):
  """
  Function to check validity of openai key
  """
  # Set the API key
  openai.api_key = openaikey
  # Test the API key by making a request to the OpenAI API
  try:
      response = openai.Model.list()
      return "Valid OpenAI API key"
  except openai.OpenAIError:    
    apikeylink = "https://beta.openai.com/account/api-keys"
    return f"Incorrect OpenAI API key provided: {openaikey}. You can find your OpenAI API key here - {apikeylink}"

def createindex(files, openaikey):
  """
  Function to create index
  """

  # Basic Checks
  if not files:
    return "Upload file before proceeding further."

  fileformatvalidity = fileformatvaliditycheck(files)

  if not fileformatvalidity:
    return "Please upload documents in pdf/txt/docx/png/jpg/jpeg format only."

  if not openaikey:
    return "Please enter your openai key."

  openaiapikeyvality = openaiapikeyvaliditycheck(openaikey)

  if openaiapikeyvality != "Valid OpenAI API key":
    return openaiapikeyvality

  # Store openai key in environment
  os.environ['OPENAI_API_KEY'] = openaikey

  # Process the Documents
  doctextlist = processfiles(files)
  documents = createdocuments(doctextlist)

  # Create index
  index = GPTListIndex(documents, chunk_size_limit = 3500)
  # Save index
  index.save_to_disk('index.json')

  return "Uploading documents successfully. OpenAI API Key provided is Valid."

def docques(query, openaikey):
  """
  Function to for quering on the index created
  """

  # Store openai key in environment
  os.environ['OPENAI_API_KEY'] = openaikey

  # Load index
  index = GPTListIndex.load_from_disk('index.json')

  # Query based on index
  response = index.query(query, response_mode="tree_summarize")

  return response

def cleartext(query, output):
  """
  Function to clear text
  """
  return ["", ""]

with gr.Blocks() as demo:
    gr.Markdown(
    """
    <h1><center><b>DocQues</center></h1>
    
    """)
    gr.Markdown(
    """
    This app answers your queries on longer and multiple documents (pdf/docx/txt/png/jpeg/jpg) you upload. It uses <a href = "https://github.com/jerryjliu/gpt_index">GPT-Index</a> and OpenAI GPT3 in the backend, get your
    <a href = "https://beta.openai.com/account/api-keys">Openai key here</a> before proceeding further.\n
    """)
    gr.Markdown(
        """
        <br>**Use this space effectively by following below 2 step process.**</br>
        *Step-1*
        <br>- Upload pdf/docx/txt/png/jpeg/jpg format documents. 
        <br>- Enter your openai key.
        <br>- Click upload and wait to see if upload is successful or not. </br>
        *Step-2*
        <br>- Enter your query. 
        <br>- Click submit.
        <br>- Check Answer </br>

        Please refer to the GitHub repo this Space is based on, here - <a href = "https://github.com/ravi03071991/DocQues">DocQues</a> .
        """
    )
    with gr.Row():
      with gr.Column():
        files = gr.File(label = "Upload pdf/docx/txt format documents.", file_count="multiple")
        openaikey = gr.Textbox(lines = 1, label = "Enter your OpenAI Key.")
        upload_button = gr.Button("Upload")
        query = gr.Textbox(lines = 2, label = "Enter Your Question.")
        submit_button = gr.Button("Submit")
      with gr.Column():
        upload_output = gr.Textbox(label = "Upload/ Error.")
        ans_output = gr.Textbox(label = "Answer.")
        clear_button = gr.Button("Clear")

    # Upload button for uploading files and openai key.
    upload_button.click(createindex, inputs=[files, openaikey], outputs= [upload_output] )

    # Submit button for submitting query.
    submit_button.click(docques, inputs=[query, openaikey], outputs= [ans_output] )

    # Clear button for clearing query and answer.
    clear_button.click(cleartext, inputs=[query, ans_output], outputs= [query, ans_output] )

demo.launch()